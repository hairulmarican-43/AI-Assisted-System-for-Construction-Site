"""
Word report generation for work-at-height compliance screening.

Uses python-docx because the report is produced at runtime inside a Streamlit
app, where Node-based generators are not available.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BAND_HEX = {
    "High": "C62828",
    "Medium": "F4A300",
    "Low": "0F9D58",
    "None": "9E9E9E",
    "-": "9E9E9E",
}

HEADER_HEX = "1F3864"
ZEBRA_HEX = "F2F2F2"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _shade(cell, hex_colour: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def _cell_text(cell, text: str, *, bold: bool = False, size: int = 9,
               colour: str | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor.from_string(colour)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


def _table(doc, headers: list[str], widths: list[float]):
    """Create a styled table with fixed column widths in inches."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (head, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        _shade(cell, HEADER_HEX)
        _cell_text(cell, head, bold=True, size=9, colour="FFFFFF")
    return table


def _add_row(table, values: list, widths: list[float],
             band: str | None = None, zebra: bool = False) -> None:
    row = table.add_row()
    for i, (val, width) in enumerate(zip(values, widths)):
        cell = row.cells[i]
        cell.width = Inches(width)
        if band and i == 0:
            _shade(cell, BAND_HEX.get(band, "9E9E9E"))
            _cell_text(cell, val, bold=True, size=9, colour="FFFFFF")
        else:
            if zebra:
                _shade(cell, ZEBRA_HEX)
            _cell_text(cell, val, size=9)


def _rule(doc) -> None:
    """Horizontal rule as a paragraph bottom border."""
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "AAAAAA")
    borders.append(bottom)
    p_pr.append(borders)
    para.paragraph_format.space_after = Pt(6)


def _footer(section, text: str) -> None:
    para = section.footer.paragraphs[0]
    para.text = text
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string("777777")


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------

def build_report(records: list, stats: dict, config: dict,
                 project: str = "", location: str = "",
                 inspector: str = "") -> bytes:
    """
    Build the .docx report and return it as bytes.

    records: list of ScreeningRecord (see app.py)
    stats:   output of compute_statistics()
    """
    doc = Document()

    # Landscape suits the wide findings tables.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height, section.page_width)
    section.left_margin = section.right_margin = Inches(0.7)
    section.top_margin = section.bottom_margin = Inches(0.6)

    _footer(
        section,
        "Preliminary visual screening — not a formal inspection and of no "
        "legal or regulatory standing under the WSH Act."
    )

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10)

    # --- Title -------------------------------------------------------------
    title = doc.add_heading("Work-at-Height Compliance Screening Report", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string(HEADER_HEX)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "SS 659 Scaffolds  ·  SS 528 Personal Fall-Arrest Systems  ·  "
        "SS 570 Anchor Devices and Horizontal Lifeline Systems"
    )
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor.from_string("555555")

    _rule(doc)

    meta = _table(doc, ["Field", "Detail"], [2.0, 7.5])
    rows = [
        ("Project", project or "—"),
        ("Location", location or "—"),
        ("Screened by", inspector or "—"),
        ("Report generated", datetime.now().strftime("%d %b %Y, %H:%M")),
        ("Images assessed", str(stats["image_count"])),
        ("Elements assessed", str(stats["element_count"])),
        ("Non-compliances found", str(stats["breach_count"])),
    ]
    for i, (k, v) in enumerate(rows):
        _add_row(meta, [k, v], [2.0, 7.5], zebra=(i % 2 == 1))

    doc.add_paragraph()

    # --- Scope and limitations --------------------------------------------
    doc.add_heading("1. Scope and limitations", level=1)
    for text in [
        "This report records a preliminary visual screening of site "
        "photographs against three Singapore Standards, carried out with "
        "computer vision assistance. Severity and likelihood are expressed on "
        "the Workplace Safety and Health Council Risk Assessment scale so that "
        "findings can be transferred directly into a risk assessment form.",

        "The screening assesses only what is visible in the supplied images. "
        "It cannot evaluate structural adequacy, load capacity, tie strength, "
        "foundation bearing, equipment inspection history, certification "
        "records, or any condition outside the camera's field of view. The "
        "absence of a finding is not evidence of compliance.",

        "This report does not constitute a formal inspection, does not "
        "discharge any duty under the Workplace Safety and Health Act or its "
        "subsidiary legislation, and does not replace examination by a "
        "competent person or, for scaffolds, an approved Scaffold Supervisor. "
        "All findings require verification on site before action is taken.",
    ]:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(8)

    # --- Risk scale --------------------------------------------------------
    doc.add_heading("2. Risk rating basis", level=1)
    doc.add_paragraph(
        "Risk Priority Number (RPN) = Severity × Likelihood, each rated 1 to 5."
    )

    scale = _table(doc, ["Rating", "Severity", "Likelihood"],
                   [1.2, 4.15, 4.15])
    sev_scale = config["severity_scale"]
    lik_scale = config["likelihood_scale"]
    for i in range(1, 6):
        _add_row(scale,
                 [str(i), sev_scale.get(i, ""), lik_scale.get(i, "")],
                 [1.2, 4.15, 4.15], zebra=(i % 2 == 0))

    doc.add_paragraph()

    bands = _table(doc, ["Band", "RPN range", "Required action"],
                   [1.2, 1.5, 6.8])
    prev_max = 0
    for band in config["risk_bands"]:
        rng = (f"{prev_max + 1}–{band['max']}"
               if prev_max else f"1–{band['max']}")
        _add_row(bands, [band["label"], rng, band["action"]],
                 [1.2, 1.5, 6.8], band=band["label"])
        prev_max = band["max"]

    # --- Statistics --------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("3. Collated severity statistics", level=1)

    doc.add_heading("3.1 By standard", level=2)
    by_std = _table(
        doc,
        ["Standard", "Elements", "Breaches", "High", "Medium", "Low",
         "Peak severity", "Peak RPN"],
        [3.0, 0.9, 0.9, 0.7, 0.85, 0.7, 1.2, 0.95],
    )
    for i, row in enumerate(stats["by_category"]):
        _add_row(by_std, [
            row["standard"], row["elements"], row["breaches"],
            row["High"], row["Medium"], row["Low"],
            row["peak_severity"] or "—", row["peak_rpn"] or "—",
        ], [3.0, 0.9, 0.9, 0.7, 0.85, 0.7, 1.2, 0.95],
            zebra=(i % 2 == 1))

    doc.add_paragraph()
    doc.add_heading("3.2 Severity distribution", level=2)
    sev_tbl = _table(
        doc, ["Severity", "Descriptor", "Breaches", "Share"],
        [1.1, 5.0, 1.2, 1.2])
    total_b = max(stats["breach_count"], 1)
    for level in range(5, 0, -1):
        count = stats["severity_distribution"].get(level, 0)
        _add_row(sev_tbl, [
            str(level), sev_scale.get(level, ""), str(count),
            f"{100 * count / total_b:.0f}%",
        ], [1.1, 5.0, 1.2, 1.2], zebra=(level % 2 == 0))

    doc.add_paragraph()
    doc.add_heading("3.3 Most frequent non-compliances", level=2)
    if stats["top_checks"]:
        top = _table(
            doc, ["Rank", "Requirement", "Standard", "Occurrences",
                  "Peak RPN"],
            [0.7, 4.0, 2.6, 1.1, 1.1])
        for i, item in enumerate(stats["top_checks"], start=1):
            _add_row(top, [
                str(i), item["check"], item["standard"],
                str(item["count"]), str(item["peak_rpn"]),
            ], [0.7, 4.0, 2.6, 1.1, 1.1], zebra=(i % 2 == 0))
    else:
        doc.add_paragraph("No non-compliances were recorded.")

    doc.add_paragraph()
    doc.add_heading("3.4 Items requiring on-site verification", level=2)
    doc.add_paragraph(
        f"{stats['not_visible_count']} check(s) could not be assessed from "
        "the images supplied. These remain open and must be verified on foot."
    )
    if stats["not_visible_by_check"]:
        nv = _table(doc, ["Requirement", "Standard", "Occurrences"],
                    [4.5, 3.5, 1.5])
        for i, item in enumerate(stats["not_visible_by_check"]):
            _add_row(nv, [item["check"], item["standard"], str(item["count"])],
                     [4.5, 3.5, 1.5], zebra=(i % 2 == 1))

    # --- Priority actions --------------------------------------------------
    doc.add_page_break()
    doc.add_heading("4. Priority actions", level=1)

    if stats["all_breaches"]:
        doc.add_paragraph(
            "Ordered by Risk Priority Number, highest first. High-band items "
            "require work to stop until controls are in place."
        )
        pri = _table(
            doc,
            ["Band", "RPN", "Source", "Element", "Requirement", "Observation",
             "Clause"],
            [0.8, 0.55, 1.5, 1.6, 1.9, 2.9, 1.3],
        )
        for item in stats["all_breaches"][:40]:
            _add_row(pri, [
                item["band"], str(item["rpn"]), item["source"],
                item["element"], item["check"], item["observation"],
                item["clause"],
            ], [0.8, 0.55, 1.5, 1.6, 1.9, 2.9, 1.3], band=item["band"])

        if len(stats["all_breaches"]) > 40:
            note = doc.add_paragraph(
                f"{len(stats['all_breaches']) - 40} further lower-priority "
                "findings are listed in section 5."
            )
            note.runs[0].font.size = Pt(8)
    else:
        doc.add_paragraph("No non-compliances were identified.")

    # --- Per-image detail --------------------------------------------------
    # No forced page break here: section 4's table often spills a row onto the
    # next page, and a break on top of that leaves a near-blank sheet.
    doc.add_paragraph()
    doc.add_heading("5. Image-by-image detail", level=1)

    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"5.{idx}  {rec.source}", level=2)

        cap = doc.add_paragraph()
        cap_run = cap.add_run(
            f"Screened {rec.timestamp:%d %b %Y, %H:%M}  ·  "
            f"{len(rec.assessment.elements)} element(s)  ·  "
            f"highest RPN {rec.assessment.max_rpn}"
        )
        cap_run.font.size = Pt(8)
        cap_run.font.color.rgb = RGBColor.from_string("666666")

        if rec.annotated_png:
            try:
                doc.add_picture(io.BytesIO(rec.annotated_png),
                                width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        if rec.assessment.truncated:
            warn = doc.add_paragraph()
            warn_run = warn.add_run(
                "Incomplete: the analysis was cut short and this image may "
                "contain further unreported elements."
            )
            warn_run.bold = True
            warn_run.font.color.rgb = RGBColor.from_string("C62828")
            warn_run.font.size = Pt(8)

        if rec.assessment.scene_notes:
            note = doc.add_paragraph(rec.assessment.scene_notes)
            note.runs[0].font.size = Pt(8)
            note.runs[0].italic = True

        det = _table(
            doc,
            ["Band", "Element", "Standard", "Requirement", "Status",
             "Observation", "S", "L", "RPN"],
            [0.8, 1.5, 1.7, 1.8, 1.0, 2.6, 0.35, 0.35, 0.45],
        )
        ordered = sorted(
            [(e, f) for e in rec.assessment.elements
             for f in e.findings],
            key=lambda pair: (-pair[1].rpn, pair[0].element_id),
        )
        for e, f in ordered:
            band = f.band if f.status == "non_compliant" else "-"
            _add_row(det, [
                band, e.element_id, f.standard_short, f.check_title,
                f.status.replace("_", " "), f.observation,
                str(f.severity), str(f.likelihood),
                str(f.rpn) if f.rpn else "—",
            ], [0.8, 1.5, 1.7, 1.8, 1.0, 2.6, 0.35, 0.35, 0.45],
                band=band if band != "-" else None)

        doc.add_paragraph()

    # --- Sign-off ----------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("6. Verification and sign-off", level=1)
    doc.add_paragraph(
        "The findings above are machine-generated and unverified. A competent "
        "person must confirm each item on site before it is closed out."
    )

    sign = _table(doc, ["", "Name", "Signature", "Date"],
                  [2.2, 3.0, 3.0, 1.8])
    for role in ["Screened by", "Verified by (competent person)",
                 "Approved by"]:
        row = sign.add_row()
        for i, width in enumerate([2.2, 3.0, 3.0, 1.8]):
            row.cells[i].width = Inches(width)
        _cell_text(row.cells[0], role, bold=True)
        for i in range(1, 4):
            _cell_text(row.cells[i], "")
        row.height = Inches(0.45)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
